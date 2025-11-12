import io
import os
from typing import Optional, Dict, Any, List
from fastapi import UploadFile, HTTPException
import boto3
from PyPDF2 import PdfReader
import docx
from sqlalchemy.orm import Session
from app.database import models
from app import schemas
from langchain.text_splitter import RecursiveCharacterTextSplitter
from app.services.llm_service import llm_service
import hashlib
from datetime import datetime
import mimetypes

# 지원되는 파일 형식
SUPPORTED_FILE_TYPES = {
    'application/pdf': '.pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
    'text/plain': '.txt',
    'text/html': '.html'
}

# 파일 크기 제한 (50MB)
MAX_FILE_SIZE = 50 * 1024 * 1024

class DocumentProcessingService:
    """
    고급 문서 처리 서비스
    파일 업로드, 텍스트 추출, 청킹, 임베딩, S3 저장을 담당
    """
    
    def __init__(self):
        """AWS S3 클라이언트 초기화"""
        try:
            self.s3_client = boto3.client(
                's3',
                region_name=os.getenv("AWS_REGION", "ap-northeast-2"),
                aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
                aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
            )
            self.bucket_name = os.getenv("S3_BUCKET_NAME", "campusmate-documents")
            print("[INFO] S3 client initialized successfully")
        except Exception as e:
            print(f"[WARNING] Failed to initialize S3 client: {e}")
            self.s3_client = None
            self.bucket_name = None
    
    def validate_file(self, file: UploadFile) -> Dict[str, Any]:
        """
        업로드된 파일의 유효성을 검사합니다.
        
        Args:
            file: 업로드된 파일
            
        Returns:
            검증 결과 딕셔너리
            
        Raises:
            HTTPException: 파일이 유효하지 않을 경우
        """
        # 파일 크기 확인
        file.file.seek(0, 2)  # 파일 끝으로 이동
        file_size = file.file.tell()
        file.file.seek(0)  # 다시 처음으로
        
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB"
            )
        
        # MIME 타입 확인
        content_type = file.content_type
        if content_type not in SUPPORTED_FILE_TYPES:
            raise HTTPException(
                status_code=415,
                detail=f"Unsupported file type: {content_type}. Supported types: {list(SUPPORTED_FILE_TYPES.keys())}"
            )
        
        # 파일 해시 생성 (중복 확인용)
        file_content = file.file.read()
        file.file.seek(0)
        file_hash = hashlib.md5(file_content).hexdigest()
        
        return {
            "file_size": file_size,
            "content_type": content_type,
            "file_hash": file_hash,
            "extension": SUPPORTED_FILE_TYPES[content_type]
        }
    
    def extract_text_from_file(self, file: UploadFile) -> Optional[str]:
        """
        파일에서 텍스트를 추출합니다.
        
        Args:
            file: 업로드된 파일
            
        Returns:
            추출된 텍스트 또는 None
        """
        print(f"--- Extracting text from {file.filename} ({file.content_type}) ---")
        
        try:
            file.file.seek(0)
            
            if file.content_type == "application/pdf":
                return self._extract_from_pdf(file)
            elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_from_docx(file)
            elif file.content_type == "text/plain":
                return self._extract_from_txt(file)
            elif file.content_type == "text/html":
                return self._extract_from_html(file)
            else:
                print(f"[ERROR] Unsupported content type: {file.content_type}")
                return None
                
        except Exception as e:
            print(f"[ERROR] Text extraction failed: {e}")
            return None
        finally:
            file.file.seek(0)
    
    def _extract_from_pdf(self, file: UploadFile) -> str:
        """PDF에서 텍스트 추출"""
        pdf_stream = io.BytesIO(file.file.read())
        reader = PdfReader(pdf_stream)
        text = ""
        
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            text += f"[페이지 {i+1}]\n{page_text}\n\n"
        
        print(f"--- Extracted {len(text)} characters from {len(reader.pages)} pages ---")
        return text.strip()
    
    def _extract_from_docx(self, file: UploadFile) -> str:
        """DOCX에서 텍스트 추출"""
        doc_stream = io.BytesIO(file.file.read())
        doc = docx.Document(doc_stream)
        text = ""
        
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # 표 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        print(f"--- Extracted {len(text)} characters from DOCX ---")
        return text.strip()
    
    def _extract_from_txt(self, file: UploadFile) -> str:
        """TXT에서 텍스트 추출"""
        content_bytes = file.file.read()
        # 인코딩 자동 감지 시도
        for encoding in ['utf-8', 'cp949', 'euc-kr']:
            try:
                text = content_bytes.decode(encoding)
                print(f"--- Decoded TXT with {encoding} encoding ---")
                return text.strip()
            except UnicodeDecodeError:
                continue
        
        # 모든 인코딩 실패시 에러 문자 무시하고 디코드
        text = content_bytes.decode('utf-8', errors='ignore')
        print(f"--- Decoded TXT with error handling ---")
        return text.strip()
    
    def _extract_from_html(self, file: UploadFile) -> str:
        """HTML에서 텍스트 추출 (BeautifulSoup 사용 권장이지만 기본 구현)"""
        content = file.file.read().decode('utf-8', errors='ignore')
        
        # 간단한 HTML 태그 제거 (실제로는 BeautifulSoup 사용 권장)
        import re
        text = re.sub(r'<[^>]+>', '', content)
        text = re.sub(r'\s+', ' ', text)  # 공백 정리
        
        print(f"--- Extracted {len(text)} characters from HTML ---")
        return text.strip()
    
    async def upload_to_s3(self, file: UploadFile, school_id: int, file_hash: str) -> Optional[str]:
        """
        파일을 S3에 업로드합니다.
        
        Args:
            file: 업로드할 파일
            school_id: 학교 ID
            file_hash: 파일 해시
            
        Returns:
            S3 URL 또는 None
        """
        if not self.s3_client:
            print("[WARNING] S3 client not available, skipping upload")
            return None
        
        try:
            # S3 키 생성: school_id/year/month/hash_filename
            now = datetime.now()
            s3_key = f"documents/{school_id}/{now.year}/{now.month:02d}/{file_hash}_{file.filename}"
            
            file.file.seek(0)
            self.s3_client.upload_fileobj(
                file.file,
                self.bucket_name,
                s3_key,
                ExtraArgs={
                    'ContentType': file.content_type,
                    'Metadata': {
                        'school_id': str(school_id),
                        'original_filename': file.filename,
                        'upload_date': now.isoformat()
                    }
                }
            )
            
            s3_url = f"https://{self.bucket_name}.s3.{os.getenv('AWS_REGION', 'ap-northeast-2')}.amazonaws.com/{s3_key}"
            print(f"--- File uploaded to S3: {s3_url} ---")
            return s3_url
            
        except Exception as e:
            print(f"[ERROR] S3 upload failed: {e}")
            return None
        finally:
            file.file.seek(0)
    
    def create_text_chunks(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """
        텍스트를 의미 있는 청크로 분할합니다.
        
        Args:
            text: 분할할 텍스트
            chunk_size: 청크 크기
            chunk_overlap: 청크 간 중복 크기
            
        Returns:
            텍스트 청크 리스트
        """
        print(f"--- Chunking text ({len(text)} chars) with size={chunk_size}, overlap={chunk_overlap} ---")
        
        if len(text.strip()) == 0:
            return []
        
        # LangChain의 RecursiveCharacterTextSplitter 사용
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )
        
        chunks = text_splitter.split_text(text)
        
        # 빈 청크 제거 및 최소 길이 필터링
        meaningful_chunks = [
            chunk.strip() for chunk in chunks 
            if chunk.strip() and len(chunk.strip()) > 50  # 최소 50자
        ]
        
        print(f"--- Created {len(meaningful_chunks)} meaningful chunks ---")
        return meaningful_chunks
    
    async def process_document(
        self, 
        school_id: int, 
        category: str, 
        file: UploadFile,
        department: Optional[str] = None,
        contact_info: Optional[str] = None,
        db: Session = None
    ) -> Dict[str, Any]:
        """
        문서를 완전히 처리합니다: 업로드, 추출, 청킹, 임베딩, DB 저장
        
        Args:
            school_id: 학교 ID
            category: 카테고리
            file: 업로드된 파일
            department: 담당 부서 (선택)
            contact_info: 연락처 (선택)
            db: 데이터베이스 세션
            
        Returns:
            처리 결과 딕셔너리
        """
        print("=== Starting Complete Document Processing ===")
        
        try:
            # 1단계: 파일 검증
            validation_info = self.validate_file(file)
            print(f"--- File validated: {validation_info['file_size']} bytes ---")
            
            # 2단계: 중복 확인 (같은 해시 값의 파일이 이미 있는지)
            if db:
                existing_doc = db.query(models.Document).filter(
                    models.Document.school_id == school_id,
                    models.Document.file_hash == validation_info['file_hash']
                ).first()
                
                if existing_doc:
                    return {
                        "status": "duplicate",
                        "message": "Document with same content already exists",
                        "existing_document_id": existing_doc.id
                    }
            
            # 3단계: 텍스트 추출
            extracted_text = self.extract_text_from_file(file)
            if not extracted_text:
                return {"status": "error", "message": "Failed to extract text from file"}
            
            # 4단계: S3 업로드 (비동기)
            s3_url = await self.upload_to_s3(file, school_id, validation_info['file_hash'])
            
            # 5단계: 데이터베이스에 문서 레코드 생성
            new_document = models.Document(
                school_id=school_id,
                category=category,
                file_name=file.filename,
                file_size=validation_info['file_size'],
                file_hash=validation_info['file_hash'],
                s3_url=s3_url,
                department=department,
                contact_info=contact_info,
                upload_date=datetime.now()
            )
            
            if db:
                db.add(new_document)
                db.commit()
                db.refresh(new_document)
                print(f"--- Document record saved with ID: {new_document.id} ---")
            
            # 6단계: 텍스트 청킹
            chunks = self.create_text_chunks(extracted_text)
            if not chunks:
                if db:
                    db.delete(new_document)
                    db.commit()
                return {"status": "error", "message": "Failed to create meaningful text chunks"}
            
            # 7단계: 임베딩 생성
            embeddings = llm_service.get_embeddings(chunks)
            if not embeddings or len(embeddings) != len(chunks):
                if db:
                    db.delete(new_document)
                    db.commit()
                return {"status": "error", "message": "Failed to generate embeddings"}
            
            # 8단계: 청크와 임베딩을 데이터베이스에 저장
            if db:
                chunk_count = 0
                for i, chunk_text in enumerate(chunks):
                    chunk_record = models.DocumentChunk(
                        document_id=new_document.id,
                        chunk_text=chunk_text,
                        chunk_index=i,
                        embedding=embeddings[i]
                    )
                    db.add(chunk_record)
                    chunk_count += 1
                
                db.commit()
                print(f"--- Saved {chunk_count} chunks with embeddings ---")
            
            print("=== Document Processing Completed Successfully ===")
            
            return {
                "status": "success",
                "message": "Document processed and stored successfully",
                "document_id": new_document.id if db else None,
                "file_name": file.filename,
                "text_length": len(extracted_text),
                "chunk_count": len(chunks),
                "s3_url": s3_url,
                "processing_details": {
                    "file_size_mb": round(validation_info['file_size'] / (1024*1024), 2),
                    "content_type": validation_info['content_type'],
                    "hash": validation_info['file_hash']
                }
            }
            
        except Exception as e:
            print(f"[ERROR] Document processing failed: {e}")
            # 실패 시 정리 작업
            if db and 'new_document' in locals():
                try:
                    db.delete(new_document)
                    db.commit()
                except:
                    pass
            
            return {
                "status": "error", 
                "message": f"Processing failed: {str(e)}"
            }
    
    async def process_text_document(
        self,
        doc: schemas.DocumentCreateFromText,
        db: Session
    ) -> Dict[str, Any]:
        """
        텍스트에서 직접 문서를 처리합니다 (RSS 크롤링 등에서 사용)
        
        Args:
            doc: 텍스트 문서 생성 스키마
            db: 데이터베이스 세션
            
        Returns:
            처리 결과 딕셔너리
        """
        print("=== Starting Text Document Processing ===")
        
        try:
            # 텍스트 해시 생성
            text_hash = hashlib.md5(doc.text.encode('utf-8')).hexdigest()
            
            # 중복 확인
            existing_doc = db.query(models.Document).filter(
                models.Document.school_id == doc.school_id,
                models.Document.file_hash == text_hash
            ).first()
            
            if existing_doc:
                return {
                    "status": "duplicate",
                    "message": "Document with same text content already exists",
                    "existing_document_id": existing_doc.id
                }
            
            # 문서 레코드 생성
            new_document = models.Document(
                school_id=doc.school_id,
                category=doc.category,
                source_url=doc.source_url,
                department=doc.department,
                contact_info=doc.contact_info,
                file_hash=text_hash,
                upload_date=datetime.now()
            )
            
            db.add(new_document)
            db.commit()
            db.refresh(new_document)
            print(f"--- Text document record saved with ID: {new_document.id} ---")
            
            # 청킹 및 임베딩 처리
            return await self._process_text_chunks(doc.text, new_document, db)
            
        except Exception as e:
            print(f"[ERROR] Text document processing failed: {e}")
            return {"status": "error", "message": f"Processing failed: {str(e)}"}
    
    async def _process_text_chunks(
        self,
        text: str,
        document: models.Document,
        db: Session
    ) -> Dict[str, Any]:
        """텍스트 청킹과 임베딩을 처리하는 공통 함수"""
        try:
            # 청킹
            chunks = self.create_text_chunks(text)
            if not chunks:
                db.delete(document)
                db.commit()
                return {"status": "error", "message": "Failed to create meaningful chunks"}
            
            # 임베딩 생성
            embeddings = llm_service.get_embeddings(chunks)
            if not embeddings or len(embeddings) != len(chunks):
                db.delete(document)
                db.commit()
                return {"status": "error", "message": "Failed to generate embeddings"}
            
            # 청크 저장
            for i, chunk_text in enumerate(chunks):
                chunk_record = models.DocumentChunk(
                    document_id=document.id,
                    chunk_text=chunk_text,
                    chunk_index=i,
                    embedding=embeddings[i]
                )
                db.add(chunk_record)
            
            db.commit()
            print(f"--- Saved {len(chunks)} text chunks with embeddings ---")
            
            return {
                "status": "success",
                "message": "Text document processed successfully",
                "document_id": document.id,
                "text_length": len(text),
                "chunk_count": len(chunks)
            }
            
        except Exception as e:
            print(f"[ERROR] Text chunk processing failed: {e}")
            db.delete(document)
            db.commit()
            return {"status": "error", "message": f"Chunk processing failed: {str(e)}"}

# 전역 문서 처리 서비스 인스턴스
document_service = DocumentProcessingService()

# 기존 함수들 유지 (하위 호환성)
def extract_text_from_file(file: UploadFile) -> Optional[str]:
    return document_service.extract_text_from_file(file)

async def process_document(school_id: int, category: str, file: UploadFile, db: Session):
    return await document_service.process_document(school_id, category, file, db=db)

async def process_text_document(doc: schemas.DocumentCreateFromText, db: Session):
    return await document_service.process_text_document(doc, db)

if __name__ == "__main__":
    print("=== Document Processing Service Test ===")
    print("This module requires file uploads and database connection for testing.")
    print("Use this module in your FastAPI app for document processing.")